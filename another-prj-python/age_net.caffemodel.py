from docx import Document

# Créer un nouveau document Word
doc = Document()

# Ajouter un titre
doc.add_heading('L’Histoire de l’Intelligence Artificielle', level=1)

# Ajouter un paragraphe d'introduction
doc.add_paragraph(
    "L'intelligence artificielle (IA) est une branche de l'informatique qui vise à créer des machines capables de simuler "
    "des comportements intelligents. Depuis ses débuts, l'IA a connu des évolutions majeures, passant de simples concepts "
    "théoriques à des technologies révolutionnaires utilisées dans presque tous les domaines de la vie."
)

# Ajouter un sous-titre
doc.add_heading('Les Débuts de l’IA', level=2)

# Ajouter un paragraphe sur les débuts
doc.add_paragraph(
    "Les bases de l’IA remontent aux années 1940 et 1950 avec l'apparition des premiers ordinateurs. Alan Turing, souvent "
    "considéré comme le père de l’informatique, a proposé le test de Turing en 1950 pour évaluer la capacité d’une machine "
    "à imiter l’intelligence humaine. En 1956, la conférence de Dartmouth marque la naissance officielle de l'IA en tant "
    "que domaine de recherche."
)

# Ajouter un sous-titre
doc.add_heading('L’Évolution de l’IA', level=2)

# Ajouter un paragraphe sur l'évolution
doc.add_paragraph(
    "Dans les années 1970 et 1980, l’IA a progressé grâce au développement des systèmes experts, qui utilisent des bases "
    "de connaissances pour résoudre des problèmes complexes. Cependant, des limites techniques et un manque de données "
    "ont ralenti les avancées. Avec l’essor des puissances de calcul et l’arrivée du big data au début des années 2000, "
    "l’apprentissage automatique et les réseaux de neurones profonds ont permis des percées majeures."
)

# Ajouter un sous-titre
doc.add_heading('L’IA Aujourd’hui', level=2)

# Ajouter un paragraphe sur l'état actuel
doc.add_paragraph(
    "Aujourd’hui, l’IA est omniprésente : des assistants virtuels comme Siri et Alexa aux voitures autonomes, en passant par "
    "les diagnostics médicaux assistés par IA. Les algorithmes de deep learning et les avancées en traitement du langage "
    "naturel ont ouvert des possibilités infinies pour l’automatisation et l’innovation."
)

# Ajouter une conclusion
doc.add_heading('Conclusion', level=2)
doc.add_paragraph(
    "L’intelligence artificielle continue de transformer notre monde à un rythme rapide. Son histoire est un témoignage "
    "de l’ingéniosité humaine, et son avenir promet encore plus de découvertes et de défis. Alors que l’IA évolue, il est "
    "essentiel de réfléchir à ses implications éthiques et sociétales pour garantir un avenir bénéfique pour tous."
)

# Enregistrer le document
doc.save('Histoire_IA.docx')

print("Le document Word a été créé avec succès !")
